#!/usr/bin/env python3
"""Build the detailed closure report and a Chinese manuscript draft.

The documents are deliberately generated only from the frozen
``paper_v24_real_01`` artifacts.  This script does not execute any experiment,
query ClickHouse, or change the v2.4 analysis outputs.
"""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RUN = ROOT / "runs" / "paper_v24_real_01"
FIG_M = RUN / "results" / "figures_manuscript" / "zh"
FIG_C = RUN / "results" / "figures_submission_core" / "zh"
OUT = ROOT / "deliverables"
OUT.mkdir(parents=True, exist_ok=True)

REPORT_PATH = OUT / "乌克兰能源冲击与互联网韧性_研究价值实验闭环与逐图论证_2026-08-08.docx"
PAPER_PATH = OUT / "乌克兰能源冲击与互联网韧性_中文论文正式草稿_2026-08-08.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4E79"
LIGHT_BLUE = "DCE6F1"
PALE_BLUE = "EEF4F8"
ORANGE = "D55E00"
GREEN = "00876C"
GRAY = "666666"
LIGHT_GRAY = "F4F6F9"
RED = "B7472A"
WHITE = "FFFFFF"
BLACK = "222222"
# The bundled headless LibreOffice renderer does not resolve Hiragino's CJK
# glyphs reliably even though macOS has the font.  Arial Unicode MS is present
# locally and renders both Chinese and Latin in Word/LibreOffice, so use one
# family for every OOXML script class to avoid tofu boxes in the deliverables.
CJK_FONT = "Arial Unicode MS"
LATIN_FONT = "Arial Unicode MS"


REFERENCES = [
    ("[1]", "Florian Holzbauer, Sebastian Strobl, Johanna Ullrich", "Tracking Internet Disruptions in Ukraine: Insights from Three Years of Active Full Block Scans", "ACM IMC 2025", "https://conferences.sigcomm.org/imc/2025/accepted-papers/"),
    ("[2]", "Florian Holzbauer, Sebastian Strobl, Johanna Ullrich", "Tracking Internet Disruptions in Ukraine（预印本）", "University of Vienna ePrints", "https://eprints.cs.univie.ac.at/8510/1/IMC25_Tracking_Internet_Disruptions_Ukraine_preprint.pdf"),
    ("[3]", "Schulman & Spring", "Pingin' in the Rain / Residential Links Under the Weather", "ACM SIGCOMM 2019", "https://conferences.sigcomm.org/sigcomm/2019/program.html"),
    ("[4]", "Marder et al.", "Access Denied: Assessing Physical Risks to Internet Access Networks", "USENIX Security 2023", "https://www.usenix.org/conference/usenixsecurity23/presentation/marder"),
    ("[5]", "Zachary S. Bischof et al.", "Destination Unreachable: Characterizing Internet Outages and Shutdowns", "ACM SIGCOMM 2023", "https://conferences.sigcomm.org/sigcomm/2023/program.html"),
    ("[6]", "Quan et al.", "Trinocular: Understanding Internet Reliability Through Adaptive Probing", "ACM SIGCOMM 2013", "https://ant.isi.edu/~johnh/PAPERS/Quan13c.html"),
    ("[7]", "Georgia Tech Internet Intelligence Lab", "IODA Glossary and Data Signals", "IODA", "https://ioda.inetintel.cc.gatech.edu/resources?tab=glossary"),
    ("[8]", "IODA", "How Russia's Recent Attacks on Ukraine's Energy Grid Impacted Internet Connectivity", "IODA report", "https://ioda.inetintel.cc.gatech.edu/reports/how-russias-recent-attacks-on-ukraines-energy-grid-impacted-its-internet-connectivity-2/"),
    ("[9]", "Ege Cem Kirci, Martin Vahlensieck, Laurent Vanbever", "Is My Internet Down? Sifting Through User-affecting Outages with Google Trends", "ACM IMC 2022", "https://conferences.sigcomm.org/imc/2022/paper-access/"),
    ("[10]", "Reethika Ramesh et al.", "Network Responses to Russia's Invasion of Ukraine in 2022", "USENIX Security 2023", "https://www.usenix.org/conference/usenixsecurity23/presentation/ramesh-network-responses"),
    ("[11]", "Yasin Alhamwy, Abdalmohimn Alshebane, Oliver Hohlfeld", "Power Out, Internet Down? Measuring Internet Resilience in South Africa", "ACM IMC 2025 Poster", "https://www.ohohlfeld.com/paper/2025-IMC-South_Africa.pdf"),
    ("[12]", "Niloofar Bayat, Kunal Mahajan, Sam Denton, Vishal Misra, Dan Rubenstein", "Down for Failure: Active Power Status Monitoring", "Future Generation Computer Systems 125 (2021), 629–640", "https://doi.org/10.1016/j.future.2021.06.055"),
    ("[13]", "Ramakrishna Padmanabhan, Aaron Schulman, Alberto Dainotti, Dave Levin, Neil Spring", "How to Find Correlated Internet Failures", "PAM 2019", "https://www.caida.org/catalog/papers/2019_how_find_correlated_internet/how_find_correlated_internet.pdf"),
    ("[14]", "Dainotti et al.", "When the Internet Sleeps: Correlating Diurnal Networks with External Factors", "ACM IMC 2014", "https://conferences.sigcomm.org/imc/2014/papers/p87.pdf"),
    ("[15]", "Emile Aben", "The Resilience of the Internet in Ukraine — One Year On", "RIPE Labs", "https://labs.ripe.net/author/emileaben/the-resilience-of-the-internet-in-ukraine-one-year-on/"),
    ("[16]", "Cristian Trusin, Leandro Bertholdo, Jose Jair Santanna", "The Effect of the Russian-Ukrainian Conflict from the Perspective of Internet eXchanges", "CNSM 2022 / arXiv", "https://arxiv.org/abs/2211.06123"),
    ("[17]", "Sangeetha Abdu Jyothi", "Characterizing the Role of Power Grids in Internet Resilience", "arXiv", "https://arxiv.org/abs/2306.02502"),
    ("[18]", "IEA", "Ukraine's Energy System Under Attack", "International Energy Agency", "https://www.iea.org/reports/ukraines-energy-security-and-the-coming-winter/ukraines-energy-system-under-attack"),
    ("[19]", "ACM SIGCOMM", "IMC 2026 Call for Papers", "Official CFP", "https://conferences.sigcomm.org/imc/2026/cfp/"),
    ("[20]", "ACM SIGCOMM", "IMC 2026 Call for Posters", "Official CFP", "https://conferences.sigcomm.org/imc/2026/cfposters/"),
    ("[21]", "PAM Steering Committee", "Passive and Active Measurement Conference", "PAM 2026 official site", "https://pam2026.at/"),
    ("[22]", "IFIP", "Traffic Measurement and Analysis Conference", "TMA 2026 official site", "https://tma.ifip.org/2026/"),
    ("[23]", "IEEE Communications Society", "IEEE Transactions on Networking", "Journal scope", "https://www.comsoc.org/publications/journals/ieee-tnet"),
    ("[24]", "Elsevier", "Computer Networks", "Journal scope", "https://www.sciencedirect.com/journal/computer-networks"),
    ("[25]", "Akshath Jain, Deepayan Patra, Peijing Xu, Justine Sherry, Phillipa Gill", "The Ukrainian Internet Under Attack: an NDT Perspective", "ACM IMC 2022", "https://doi.org/10.1145/3517745.3561449"),
    ("[26]", "British Journal of Political Science authors", "Strategic Interdependence: Using Internet Outage Data to Study How Combatants Manage Collective Institutions During War", "BJPolS 2026", "https://www.cambridge.org/core/journals/british-journal-of-political-science/article/strategic-interdependence-using-internet-outage-data-to-study-how-combatants-manage-collective-institutions-during-war/78A03E467A5BD56CD92D213D32AD6617"),
]


class Numbering:
    """Create actual Word list numbering rather than typed glyphs."""

    def __init__(self, doc: Document):
        self.part = doc.part.numbering_part.element
        self.next_abs = 100
        self.next_num = 100
        self.bullet = self._new("bullet")
        self.decimal = self._new("decimal")

    def _new(self, kind: str) -> int:
        abstract_id = self.next_abs
        self.next_abs += 1
        num_id = self.next_num
        self.next_num += 1
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        for level in range(3):
            lvl = OxmlElement("w:lvl")
            lvl.set(qn("w:ilvl"), str(level))
            start = OxmlElement("w:start")
            start.set(qn("w:val"), "1")
            lvl.append(start)
            num_fmt = OxmlElement("w:numFmt")
            num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
            lvl.append(num_fmt)
            lvl_text = OxmlElement("w:lvlText")
            if kind == "bullet":
                lvl_text.set(qn("w:val"), ["•", "–", "◦"][level])
            else:
                lvl_text.set(qn("w:val"), "%" + str(level + 1) + ".")
            lvl.append(lvl_text)
            suff = OxmlElement("w:suff")
            suff.set(qn("w:val"), "tab")
            lvl.append(suff)
            ppr = OxmlElement("w:pPr")
            tabs = OxmlElement("w:tabs")
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "num")
            tab.set(qn("w:pos"), str(540 + level * 360))
            tabs.append(tab)
            ppr.append(tabs)
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), str(540 + level * 360))
            ind.set(qn("w:hanging"), "260")
            ppr.append(ind)
            lvl.append(ppr)
            abstract.append(lvl)
        self.part.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_ref = OxmlElement("w:abstractNumId")
        abs_ref.set(qn("w:val"), str(abstract_id))
        num.append(abs_ref)
        self.part.append(num)
        return num_id

    def apply(self, paragraph, kind: str = "bullet", level: int = 0):
        ppr = paragraph._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), str(level))
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), str(self.bullet if kind == "bullet" else self.decimal))
        num_pr.append(ilvl)
        num_pr.append(num_id)
        ppr.append(num_pr)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_on_open(doc: Document):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def set_run_font(run, latin=LATIN_FONT, east_asia=CJK_FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])
    return run


def add_hyperlink(paragraph, text: str, url: str):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)
    rpr.extend([rfonts, color, underline])
    run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_document(doc: Document, short_title: str):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    specs = {
        "Title": (28, DARK_BLUE, True, 0, 14),
        "Subtitle": (13, GRAY, False, 0, 8),
        "Heading 1": (16, BLUE, True, 18, 10),
        "Heading 2": (13, BLUE, True, 12, 6),
        "Heading 3": (12, DARK_BLUE, True, 8, 4),
        "Caption": (9, GRAY, False, 4, 8),
    }
    for name, (size, color, bold, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = LATIN_FONT
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = name.startswith("Heading")
    doc.styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.styles["Subtitle"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.styles["Caption"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for style_name, size, color, fill in (
        ("Quote", 10.5, DARK_BLUE, None),
        ("Intense Quote", 10.5, DARK_BLUE, PALE_BLUE),
    ):
        style = doc.styles[style_name]
        style.font.name = LATIN_FONT
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.paragraph_format.left_indent = Inches(0.25)
        style.paragraph_format.right_indent = Inches(0.20)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(8)

    header = sec.header.paragraphs[0]
    header.text = short_title
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=8.5, color=GRAY)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("第 ")
    set_run_font(r, size=8.5, color=GRAY)
    add_field(footer, "PAGE")
    r = footer.add_run(" 页")
    set_run_font(r, size=8.5, color=GRAY)
    set_repeat_on_open(doc)


def add_cover(doc: Document, title: str, subtitle: str, doc_type: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(doc_type.upper())
    set_run_font(r, size=10, color=ORANGE, bold=True)
    p = doc.add_paragraph(style="Title")
    p.add_run(title)
    p = doc.add_paragraph(style="Subtitle")
    p.add_run(subtitle)
    doc.add_paragraph("")
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    t.columns[0].width = Inches(6.35)
    c = t.cell(0, 0)
    set_cell_shading(c, DARK_BLUE)
    set_cell_margins(c, top=180, bottom=180, start=220, end=220)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("冻结运行：paper_v24_real_01  ·  文档日期：2026-08-08\n")
    set_run_font(r, size=10.5, color=WHITE, bold=True)
    r = p.add_run("结论状态：科学闭环已完成（负向/边界型）；投稿证据包尚需归档性补齐")
    set_run_font(r, size=10.5, color=WHITE)
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run("核心问题")
    set_run_font(r, size=11, color=ORANGE, bold=True)
    p = doc.add_paragraph(style="Intense Quote")
    p.add_run("精确登记的计划停电能否作为弱监督，校准一组对供电变化敏感的互联网端点；冻结这组端点后，能否量化留出的战时能源攻击，并判断 ASN×一级行政区响应是否形成可重复、可预测的韧性指纹？")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("内部研究与投稿准备材料")
    set_run_font(r, size=10, color=GRAY)
    p = doc.add_paragraph()
    r = p.add_run("注意：本文件不重跑 v2.4，也不根据结果重新选择阈值、窗口、事件或模型。")
    set_run_font(r, size=9.5, color=GRAY, italic=True)
    doc.add_page_break()


def add_toc(doc: Document, entries):
    doc.add_heading("目录", level=1)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u')
    p = doc.add_paragraph()
    r = p.add_run("如目录页未自动显示，请在 Word 中右键目录并选择“更新域”。")
    set_run_font(r, size=9, color=GRAY, italic=True)
    p = doc.add_paragraph()
    r = p.add_run("章节导航")
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    for item in entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run_font(r, size=9.5, color=BLACK)
    doc.add_page_break()


def add_para(doc: Document, text: str, *, bold_lead: str | None = None, style=None, align=None):
    p = doc.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, bold=True)
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    if align is not None:
        p.alignment = align
    return p


def add_bullets(doc: Document, numbering: Numbering, items, *, level=0, kind="bullet"):
    for item in items:
        p = doc.add_paragraph()
        numbering.apply(p, kind=kind, level=level)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_callout(doc: Document, title: str, body: str, tone: str = "blue"):
    fill, accent = (PALE_BLUE, BLUE) if tone == "blue" else ("FFF2CC", ORANGE) if tone == "orange" else ("FDE9E7", RED)
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Inches(6.5)
    c = t.cell(0, 0)
    set_cell_shading(c, fill)
    set_cell_margins(c, top=140, bottom=140, start=180, end=180)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title + "\n")
    set_run_font(r, size=10.5, color=accent, bold=True)
    r = p.add_run(body)
    set_run_font(r, size=10.2, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers, rows, widths=None, font_size=8.8, caption=None):
    if caption:
        p = doc.add_paragraph(style="Caption")
        p.paragraph_format.keep_with_next = True
        p.add_run(caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, h in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        if widths:
            cell.width = Inches(widths[idx])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(str(h))
        set_run_font(r, size=font_size, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if widths:
                cell.width = Inches(widths[idx])
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_equation(doc: Document, text: str, explanation: str | None = None):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    set_cell_shading(c, "F8F8F8")
    set_cell_margins(c, top=120, bottom=120, start=180, end=180)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, latin="Menlo", east_asia=CJK_FONT, size=10.2, color=DARK_BLUE)
    if explanation:
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2.paragraph_format.space_after = Pt(0)
        r = p2.add_run(explanation)
        set_run_font(r, size=9.2, color=GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc: Document, path: Path, caption: str, width=6.55):
    if not path.exists():
        add_callout(doc, "图像缺失", f"未找到：{path}", tone="red")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", caption.split("。", 1)[0])
    c = doc.add_paragraph(style="Caption")
    c.add_run(caption)


def add_reference_list(doc: Document):
    doc.add_heading("参考文献与官方来源", level=1)
    for label, author, title, venue, url in REFERENCES:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(-0.28)
        p.paragraph_format.left_indent = Inches(0.28)
        r = p.add_run(f"{label} {author}. {title}. {venue}. ")
        set_run_font(r, size=9.5)
        add_hyperlink(p, "在线来源", url)


def add_report_figure_card(doc: Document, number: str, title: str, image_path: Path, fields: list[tuple[str, str]], caption: str):
    doc.add_heading(f"{number}　{title}", level=2)
    add_figure(doc, image_path, caption)
    rows = [(k, v) for k, v in fields]
    add_table(doc, ["读图问题", "具体解释"], rows, widths=[1.25, 5.05], font_size=9.0)


def build_report():
    doc = Document()
    configure_document(doc, "乌克兰能源冲击与互联网韧性｜研究闭环报告")
    doc.core_properties.title = "乌克兰能源冲击与互联网韧性：研究价值、实验闭环与逐图论证"
    doc.core_properties.subject = "v2.4 冻结运行的科学审计与投稿论证"
    doc.core_properties.author = "研究团队"
    add_cover(
        doc,
        "乌克兰能源冲击与互联网韧性：\n研究价值、实验闭环与逐图论证",
        "为什么值得做、为什么这样设计、结果究竟支持什么，以及如何让审稿人相信",
        "Research closure dossier",
    )
    add_toc(doc, [
        "执行摘要：先给结论",
        "第一部分　为什么值得做：研究问题、文献位置与不可替代性",
        "第二部分　为什么这样设计实验：数据、对照、指标与因果边界",
        "第三部分　论文图像如何讲清核心结论",
        "第四部分　投稿定位、导师答辩与执行路线",
        "附录：关键文件、最终结论与参考文献",
    ])
    nums = Numbering(doc)

    doc.add_heading("执行摘要：先给结论", level=1)
    add_callout(
        doc,
        "最终判定",
        "这项研究值得做，方向没有跑偏；但原先设想的“计划停电→供电敏感 B2 传感器→攻击泛化→稳定 ASN×地区指纹”正向链条没有成立。当前 v2.4 已经形成一种更可信、更可发表的负向/边界型闭环：长期主动测量确实能量化部分攻击时点附近的互联网可达性下降和恢复差异，但 L3 国家×时间的计划停电弱标签不足以稳定校准优于 B1 稳定端点的 B2 传感器，ASN×Admin1 响应也未形成稳定可预测的固定指纹。",
    )
    add_para(doc, "这意味着“能否收尾”的答案必须分成两层。科学上可以收尾，因为每一条主假设都经过了预先冻结的检验，并允许失败；工程上也不应重跑 v2.4 或事后调参。材料发布层面还不完全收尾，因为连续 ERA5 天气协变量、Telegram 原生 JSON 与个别网页原始响应尚未全部归档。它们用于证据可追溯性和混杂敏感性，不应被用来改变主结论。")
    add_table(doc, ["核心命题", "冻结结果", "允许写入论文的结论"], [
        ("计划停电可校准供电敏感端点", "B2−B1 ΔAUPRC=+0.007885；95% CI [−0.025011,+0.037959]；置换 p=0.323", "不支持。只能说存在方向一致但事件依赖的弱信号。"),
        ("校准后传感器可泛化到留出攻击", "B2 的最大缺口仅在 5 个可推断事件中的 3 个大于 B1；AUC 仅 2/5 更大", "不支持稳定泛化。主分析必须使用冻结回退的 B1。"),
        ("攻击冲击可被主动测量量化", "6 个攻击事件均形成曲线；5 个逐事件通过推断门，Sumy 因预趋势失败仅描述；预注册主要攻击子集有 3 个可推断", "部分支持。可量化部分事件的可观测网络影响，但不是电网实际停电率。"),
        ("ASN×Admin1 构成韧性指纹", "主重复性 ρ=0.179，CI 下界跨 0，ICC=0；预测表观改善 41.7%，置换 p=0.294", "不支持固定、可预测指纹；事件与残差变化占主导。"),
        ("次要机制", "恢复债务对累计缺口有条件关联；52 个高质量同目标路径单元中 6 个 ASGeo 变化过 FDR", "仅作为条件性证据，不能挽救主链。"),
    ], widths=[1.65, 2.45, 2.25], font_size=8.6, caption="表 1　当前研究闭环的一页式判定")

    doc.add_heading("第一部分　为什么值得做：研究问题、文献位置与不可替代性", level=1)
    doc.add_heading("1. 核心问题究竟是什么", level=2)
    add_para(doc, "论文不是泛泛地问“战争会不会导致互联网中断”，也不是再做一篇事件时间线描述。最精确的问题是：外部世界中能够事先获知但并不等同于端点实际断电的计划停电记录，能否作为弱监督，从长期 IP 级主动测量里筛出一种可迁移的供电敏感传感器；若可以，这些传感器是否能在没有参与筛选的战争能源攻击上产生更强、更早、更可解释的网络信号；最后，这些冲击响应是否稳定地属于某个 ASN×一级行政区，而不是只属于某一次事件。")
    add_equation(doc, "计划停电弱标签 → 冻结端点传感器 → 留出攻击 → ASN×Admin1 重复性/预测", "任何箭头失败，都必须保留为结论；后面的次要证据不能反向替代前面的校准门。")
    add_para(doc, "三个可证伪子问题分别对应三个门：RQ1 是校准门，检验 B2 是否显著且稳健地优于 B1；RQ2 是迁移门，检验 B2 是否在留出攻击上持续增强最大缺口和累计缺口；RQ3 是指纹门，检验跨事件排序相关、ICC 与严格滚动留出预测。这个结构之所以科学，是因为它把“看到了波动”提升成“一个可被证伪的测量仪器是否成立”。")

    doc.add_heading("2. 方向有没有跑偏", level=2)
    add_para(doc, "没有跑偏。所有主实验仍围绕“能源冲击如何在互联网主动测量中可见，以及这种可见性是否能够由计划停电校准并跨事件复用”。实验 D 的恢复债务和实验 E 的 AS/ASGeo 变化只有在这个主问题下才有意义：它们描述冲击后的持续性与存活转发路径，而不是另起一篇路由论文。IODA、Cloudflare 式外部数据和运营商记录也只承担三角核验、暴露登记或证伪，不参与挑选 B2 阈值。")
    add_callout(doc, "防跑偏规则", "如果一项分析不能回答“标签是否可信、传感器是否泛化、冲击是否可量化、指纹是否重复/可预测、恢复与路径证据的边界是什么”之一，就不应进入正文。", tone="orange")

    doc.add_heading("3. 别人已经做过什么", level=2)
    add_para(doc, "相关工作并非空白。真正有说服力的写法不是声称“前人都没做”，而是准确划分前人解决了哪一段、我们验证了哪一条尚未闭合的链。最接近的研究至少有五类。")
    add_table(doc, ["研究类别/代表", "已经解决的问题", "仍未解决或与本文不同之处", "本文如何承接"], [
        ("乌克兰三年全地址块主动扫描，IMC 2025 [1][2]", "两小时扫描揭示战争期间地区/AS/地址块中断，并把网络异常与能源冲击关联。", "没有把计划停电当作弱监督去筛 IP 传感器，也没有验证 scheduled→attack 的迁移与 ASN×Admin1 指纹预测。", "必须承认它是最近邻工作；本文的新增点是“校准是否成立”及其失败边界，而不是“首次测量乌克兰”。"),
        ("Down for Failure [12]", "以住宅 IP 动态可信度和公用事业停电数据检测飓风停电。", "依赖更接近真实客户停电率的县/州数据；不是战争环境、L3 调度弱标签、留出攻击迁移或固定指纹。", "直接对比监督粒度：精细真实停电数据可以工作，并不保证粗计划表可以工作。"),
        ("南非 planned vs unplanned load shedding，IMC 2025 poster [11]", "比较计划削减与非计划削减对 NDT/BGP 的影响，发现计划削减信号较弱。", "没有构造 IP watchlist、没有冻结后对攻击做事件级迁移，也没有恢复/指纹预测。", "其发现与本文负向校准相互印证，说明计划负荷管理可能被备用电与通知行为稀释。"),
        ("Trinocular、相关故障检测、天气与日周期 [3][6][13][14]", "建立主动探测、相关故障聚合、外部因素控制与可靠端点选择方法。", "通常不面对战时多重混杂、计划表版本、攻击/停电/网络三时钟和行政区暴露。", "本文把经典方法放进更难的弱监督、战争事件和跨事件泛化设置。"),
        ("IODA、RIPE Atlas、IXP/NDT/九源战争研究 [7][8][10][15][16][25][26]", "从 BGP、活跃探测、暗网、流量、路径或州—日中断展示乌克兰网络韧性、异常与战争过程。", "多为宏观异常、路由/流量变化、政治行为解释或事后事件关联，不回答计划停电能否训练微观端点传感器。", "作为独立外部参照和背景，而不是本文标签或结果选择器。"),
        ("物理风险与电网—互联网依赖 [4][17]", "研究互联网接入设施的物理暴露和电网依赖模型。", "多依赖设施/拓扑建模或风险地图，不是从 3M IP 的长期响应中实证校准。", "说明问题的重要性，并把本文定位为观测层面的实证压力测试。"),
    ], widths=[1.45, 1.65, 1.75, 1.55], font_size=8.0, caption="表 2　与最接近工作的差异不是“有没有测网络”，而是“是否完成弱监督迁移链”")

    doc.add_heading("4. 真正“别人难做、我们能做”的点", level=2)
    add_para(doc, "不能把独特性写成“网络测量就是这样做，所以数据大就是贡献”。真正的稀缺性是四种条件同时存在，而且已经在同一冻结分析中连起来：")
    add_bullets(doc, nums, [
        "长期、固定节奏、目标集合可追踪的主动测量：2024-06-22 08:00 UTC 至 2025-01-09 14:00 UTC，2 小时一轮；不是事件发生后才开始抓取。",
        "端点级分母契约：3,042,264 个目标 IP；完整扫描周期里缺行被定义为非响应，采集文件缺失则定义为 acquisition gap，避免把“没拿到数据”误写成“端点掉线”。",
        "计划停电的版本化证据登记：本地 Europe/Kyiv 时间、UTC、队列数、零队列间隙、取消/恢复/覆盖关系、证据 URL 与哈希。它允许把一个模糊新闻事件转成可审计弱标签。",
        "真正的时间外推：训练计划停电只构造端点分数；留出计划事件检验 AUPRC；一旦门失败即冻结回退到 B1；攻击、预测和外部核验均不能反过来调 B2。",
        "三时钟分离：attack_start_utc、outage_start_utc、network_anomaly_start_utc 不合并；第三方网络异常只能做 replication anchor，不能定义电力处理组。",
        "同一框架同时给出正证据和失败边界：它不仅展示哪些攻击可见，还揭示为什么粗调度标签不能自动升级为 IP 级电力真值。",
    ])
    add_callout(doc, "独特性应怎样表述", "不是“只有我们拥有网络探测”，也不是“首次发现战争影响网络”。更准确的是：我们拥有一条可事前冻结、可逐箭头证伪的 scheduled-outage weak supervision → held-out wartime attacks → group fingerprint 测量链，并用长期端点级分母、版本化能源事件证据和严格时间外推检验了它的适用边界。")

    doc.add_heading("5. 这项研究的价值足以支撑论文吗", level=2)
    add_para(doc, "足以支撑一篇严谨的互联网测量论文，但稿件定位必须从“成功构建电力传感器系统”改成“对一种有吸引力但危险的弱监督范式进行大规模、真实战争环境的压力测试”。测量论文的价值不要求每个假设为正；它要求问题重要、数据稀缺、设计可证伪、负结果能阻止社区采用错误做法，并且给出下一步需要什么粒度的真值。")
    add_para(doc, "本文最有价值的结论是一个边界：精确到日期、时段和队列数的计划停电仍然不等于端点实际断电。用户提前避险、基站/ISP 备用电源、队列覆盖不均、调度临时取消、热负荷、战争损伤与网络自身故障会共同稀释 IP 响应。因此，若没有地址/馈线/队列映射或足够多的独立执行事件，计划表不能被直接当作可迁移的端点监督标签。这个结果会改变其他研究者的实验设计。")

    doc.add_heading("6. 已经得到的核心结论", level=2)
    add_bullets(doc, nums, [
        "数据可用性成立：观测支持内 2,416 个周期中 2,322 个完整，完整率 96.11%；乌克兰分析人口有效 ASN 比例 98.17%；端点级 Admin1×ASN 可用于地区分析的 IP 比例 68.79%。",
        "稳定端点筛选有价值：B0 AUPRC=0.207，B1=0.300，说明先排除不稳定端点能显著提高检测质量。",
        "供电特异筛选未被验证：B2 AUPRC=0.308，仅比 B1 高 0.007885，置信区间跨 0，置换 p=0.323。B2 只有 24,698 个 IP，而 B1 有 361,871 个，样本压缩没有换来可靠泛化。",
        "计划事件异质性是关键发现：7 月 28 日的 ΔAUPRC=+0.3937，而 8 月 19/20/21 日分别只有 +0.0531、+0.0081、+0.0060。移除 late-July cluster 后事件等权增益仅 +0.0224，说明效果被少数环境支配。",
        "攻击影响可以被 B1 描述：8 月 26 日最大缺口 0.1733、t90=100h；11 月 17 日 0.0262、64h；11 月 28 日 0.0282、2h；12 月 13 日 0.1177、10h；12 月 25 日 0.0102、66h。Sumy 9 月 17 日表面缺口 0.1696，但预趋势等价失败，仅作描述。",
        "B2 对留出攻击没有一致优势：在 5 个可推断事件中，B2 最大缺口 3/5 更大、累计缺口仅 2/5 更大；这不是可重复迁移。",
        "固定指纹不成立：跨事件累计缺口主 Spearman ρ=0.179，置信区间下界 −0.0115；ICC=0。M4 历史 Ridge 的 MAE=1.492，M3 group baseline=2.562，表观改善 41.7%，但置换 p=0.294。",
        "事件因素与未解释变化主导：最大缺口方差中事件约 28.9%，ASN 6.9%，Admin1 1.9%，ASN×Admin1 4.4%，残差 57.9%；累计缺口残差 64.9%；t90 残差 86.2%。",
        "恢复债务和路径变化是条件性信号：预事件网络债务与后续累计缺口有关，但预注册官方 168h 暴露缺少事件内变异；路径分析只有 52/1,250 个组—事件满足高质量门，6 个 ASGeo 单元过 FDR。",
    ])

    doc.add_heading("7. 什么叫“完美闭环”——不是每个结果都显著", level=2)
    add_para(doc, "科学闭环的定义应是：研究问题被清楚提出；监督与结果来源在时间上独立；可证伪门槛在看结果前冻结；每个失败都有明确的回退策略；最终声明严格受门槛约束；原始事件证据和计算产物可追溯。按这个定义，v2.4 已经完成闭环。按“预想的正向结论必须全部显著”的定义，它没有完成，也不应通过事后调整制造完成。")
    add_table(doc, ["闭环环节", "预期", "实际", "是否闭合"], [
        ("数据分母", "能区分非响应与采集缺口", "完整周期缺行=零响应；缺文件=不支持周期", "是"),
        ("事件登记", "电力暴露独立于网络结果", "三时钟、证据来源、窗口与版本冻结", "是；归档材料待补"),
        ("计划停电校准", "B2 显著优于 B1", "未过 CI/置换门", "是，负向"),
        ("传感器冻结", "门通过用 B2，否则回退 B1", "主方法冻结为 B1", "是"),
        ("留出攻击", "观察攻击影响并检查 B2 迁移", "B1 可量化部分事件；B2 不一致", "是，混合"),
        ("指纹", "同组可重复且可预测", "重复性与预测门失败", "是，负向"),
        ("机制证据", "恢复债务/路径适应提供解释", "仅部分、条件性", "是，但不升级主声明"),
        ("发布证据", "全部来源快照与协变量可复现", "Telegram JSON、ERA5 与少数原始页待补", "未完全"),
    ], widths=[1.15, 1.7, 2.45, 1.0], font_size=8.4, caption="表 3　科学闭环与投稿材料闭环必须分开判断")

    doc.add_heading("8. 如何让审稿人相信", level=2)
    add_bullets(doc, nums, [
        "把主张写窄：可达性缺口是“observable network reachability deficit”，不是停电户数、供电中断率或因果停电概率。",
        "把数据契约写在方法前：目标集合、扫描间隔、完整周期定义、缺行语义、采集边界、U2/U3 地理筛选必须可复现。",
        "公开事件证据账本：每条记录给出来源角色、原始时区、UTC 转换、版本、覆盖/取消、抓取时间、快照路径与 SHA-256。",
        "显式展示失败：校准 CI 跨 0、Sumy 预趋势失败、falsification 不可估计、IODA 仅 1/6 在 6 小时内一致、ICC=0，全部进入图或表。",
        "避免数据循环：网络异常来源不定义电力处理组；IODA 不挑事件；攻击结果不调 B2；测试事件不生成自身历史特征。",
        "分清主结果与次要结果：恢复债务、路径和外部平台都不能“救活”校准与指纹主假设。",
        "提供一条命令复现文档中的表和图；但 v2.4 冻结产物不再重跑。文稿中记录 commit/hash、环境、随机种子与所有门槛。",
    ])

    doc.add_heading("第二部分　为什么这样设计实验：数据、对照、指标与因果边界", level=1)
    doc.add_heading("9. 数据现状与主动测量优势", level=2)
    add_table(doc, ["维度", "冻结数据", "研究意义"], [
        ("观测范围", "2024-06-22 08:00 UTC—2025-01-09 14:00 UTC", "覆盖夏季计划停电与多轮秋冬能源攻击，支持时间外推。"),
        ("测量频率", "2 小时一轮，Frankfurt 单源主动探测", "比日级平台更接近停电窗口，但仍无法定位 2 小时内细节。"),
        ("目标规模", "3,042,264 IP；28,197 个 /24 映射", "允许先在 IP 层学稳定性，再在 /24×ASN×Admin1 聚合减少相关性。"),
        ("U2", "2,145,724 个 Ukraine+valid ASN IP；24,081 /24；1,542 ASN", "国家分析主宇宙。"),
        ("U3", "2,092,860 个 Ukraine+ASN+Admin1 IP；23,307 /24；1,536 ASN", "地区分析宇宙；排除 country-only/unknown。"),
        ("支持完整性", "2,416 支持周期；2,322 完整；96.11%", "核心质量门通过。"),
        ("传感器", "B1 361,871 IP；B2 24,698 IP；地区 B1 346,522、B2 24,236", "B2 是嵌套候选，不是独立样本；主结果因门失败使用 B1。"),
    ], widths=[1.15, 2.05, 3.15], font_size=8.6, caption="表 4　主动测量数据的规模、可用范围和解释边界")
    add_para(doc, "主动测量的最大优势不是“ping 很多”，而是同一批目标在事件之前、之中、之后持续接受同一协议与同一节奏的观察。它使每个端点能够贡献自己的历史响应概率 pN，从而用 expected_response_n=ΣpN 形成期望分母；这样，一个周期中缺失的响应可以被量化，而不是只对返回行求平均。")
    add_equation(doc, "normalized_reach(g,t) = responders(g,t) / Σᵢ∈g pNᵢ", "若只在返回行中求平均，停电造成的未返回端点会从数据表消失，指标反而看不到中断。")
    add_para(doc, "局限同样必须写清：单一 Frankfurt 视角把国际路径、源端拥塞和区域末端故障混合在一起；Ping 不响应可能来自防火墙、主机睡眠、ISP 策略或链路故障；主动测量只能得到“网络可观测性”，不能直接观测电池、基站供电或居民实际断电。")

    doc.add_heading("10. 多元情报如何确定事件时间点", level=2)
    add_para(doc, "这里的“多元情报”应在论文中称为“triangulated evidence registry（多源证据三角登记）”，避免让人误解为不可审计的情报判断。每个事件至少区分电力暴露来源、攻击事实来源和第三方网络观测来源。官方能源/运营商记录定义计划或断电窗口；政府、联合国、人道报告或可靠新闻确定攻击时间与影响范围；IODA/Cloudflare/RIPE 只验证网络异常是否在相近时间和空间出现。")
    add_table(doc, ["字段", "来源优先级", "在模型中的角色", "禁止用途"], [
        ("attack_start_utc", "官方空袭/政府/联合国/可信报道", "最早可信处理边界；保护清洁基线", "不能由自有网络曲线反推"),
        ("outage_start_utc", "Ukrenergo、州级 DSO、权威停电记录", "主 confirmatory power anchor", "不能把未映射队列直接赋给 IP"),
        ("network_anomaly_start_utc", "IODA/Cloudflare/RIPE 等第三方", "独立 replication anchor", "不能定义处理组、调阈值或选择事件"),
        ("schedule_version / supersedes", "最终调度、取消、恢复帖", "生成每 2 小时队列剂量与零队列窗口", "不能混用早期计划和最终执行状态"),
        ("source_sha256", "本地原始快照/Telegram JSON", "证明当时使用的证据未改变", "网页 URL 本身不能替代不可变快照"),
    ], widths=[1.35, 1.45, 1.75, 1.75], font_size=8.3, caption="表 5　三时钟和证据角色是避免循环论证的关键")
    add_para(doc, "清洁基线从最早可能受影响的时间再向前留出 6 小时缓冲，并回看 168 小时。若攻击早于停电，则攻击至停电的区间被标记为 transition，而不是伪装成 untreated pretrend。这一设计专门避免“用已受攻击的数据证明攻击后下降”。")

    doc.add_heading("11. 端点校准：为什么有 B0、B1、B2", level=2)
    add_para(doc, "B0 是所有在训练正常或计划窗口中至少响应过一次的端点；B1 要求正常期后验响应概率 pN 达到稳定阈值且暴露周期足够；B2 在 B1 内进一步要求计划停电期响应概率 pP 相对 pN 的下降，其 95% 保守下界 S_lo 大于 0。用 Jeffreys Beta(0.5,0.5) 后验而非裸比例，是为了避免低暴露端点因一次偶然不响应获得极端分数。")
    add_equation(doc, "pN = E[Beta(xN+0.5, nN−xN+0.5)]；pP 类似；S = pN−pP；B2 ⇔ B1 ∧ S_lo>0", "B2 的语义是“在训练计划窗口中表现出可信响应下降的稳定端点”，不是“已经确认接在某个停电馈线上”。")
    add_para(doc, "训练只使用 2024-07-07 和 2024-07-20。正常对照按星期×2 小时时间槽匹配，且严格早于第一个留出计划停电锚点。留出验证包括 07-28、08-19、08-20、08-21 和发布时排除的 12-09 混杂事件。AUPRC 适合此处，因为正类计划停电周期占比约 22.6%，类别不平衡，ROC-AUC 会掩盖实际精确率。")
    add_equation(doc, "校准主效应 = AUPRC(B2) − AUPRC(B1)", "成功门不仅要求点估计为正，还要求事件/地区块 bootstrap 的 95% CI 下界>0，并通过事件置换；同时要求至少两个独立发布集群。")

    doc.add_heading("12. 实验 A：计划停电校准的具体设计", level=2)
    add_table(doc, ["设计要素", "具体做法", "为什么科学"], [
        ("标签", "每个完整 2h 周期按最终计划表队列剂量标为 positive/registered zero；L3 国家×时间弱监督", "不伪造 IP 级真值；保留标签质量等级。"),
        ("对照", "同星期×同 2h slot，优先同日登记零队列，其次历史清洁周期，1:4 匹配", "控制昼夜、星期和长期响应习惯。"),
        ("评分", "每周期对 B0/B1/B2 计算平均响应缺口作为分类分数", "直接比较筛选带来的增量，而不是只看 B2 自身表现。"),
        ("统计", "全 PR 曲线、AUPRC、事件/Admin1 块 bootstrap、事件置换、事件级 Δ、leave-cluster-out", "承认同一事件和地区内相关性，检验是否被单次事件支配。"),
        ("回退", "若 B2 未过门，后续主方法自动冻结为 B1", "防止失败后仍把 B2 当作成功传感器。"),
    ], widths=[1.05, 2.55, 2.65], font_size=8.5, caption="表 6　实验 A 的识别目标是“增量校准”，不是“停电日是否有下降”")
    add_para(doc, "实际结果说明 B1 的稳定性筛选确实有用，但计划停电特异筛选没有提供可靠增量。07-28 几乎被 B2 完美区分，而 8 月三次只略有改善；这提示同一计划表的实际执行、备用电、温度负荷和端点组成差异远大于一个固定 IP 属性。")

    doc.add_heading("13. 实验 B：留出攻击的量化与对照", level=2)
    add_para(doc, "实验 B 同时回答两个不同问题：第一，冻结的主传感器面板能否在预先登记的攻击时点观察到网络可达性缺口；第二，若 B2 真的是供电特异传感器，它是否比 B1 在这些完全留出的攻击中给出更强信号。由于校准门失败，第一问的主分析使用 B1；第二问只作为 B2-vs-B1 的方法敏感性，不重新定义主结果。")
    add_table(doc, ["事件", "范围/角色", "主对照", "最大缺口", "AUC(h)", "t90", "推断状态"], [
        ("2024-08-26", "全国主要攻击", "清洁同时间槽自基线", "0.1733", "2.834", "100h", "通过"),
        ("2024-09-17 Sumy", "盲测/区域", "同 ASN 最近 /24", "0.1696", "2.546", "4h", "预趋势失败，仅描述"),
        ("2024-11-17", "区域主要攻击", "同 ASN 未暴露地区", "0.0262", "0.828", "64h", "通过"),
        ("2024-11-28", "区域主要攻击", "同 ASN 未暴露地区", "0.0282", "0.259", "2h", "通过"),
        ("2024-12-13", "全国压力测试", "清洁同时间槽自基线", "0.1177", "0.618", "10h", "通过；天气混杂"),
        ("2024-12-25", "全国压力测试", "清洁同时间槽自基线", "0.0102", "0.348", "66h", "通过；节假日混杂"),
    ], widths=[1.1, 1.25, 1.6, .75, .65, .55, 1.0], font_size=7.9, caption="表 7　实验 B 冻结主结果；缺口是可观测网络指标，不是供电中断率")
    add_para(doc, "全国事件没有自然的未暴露乌克兰地区，因此用 168 小时清洁期内同星期×时间槽的自身基线估计偏离；区域事件使用同 ASN、不同 Admin1 的最近 /24 匹配，以减少 ASN 技术栈差异。区域设计要求至少 30 对且匹配后所有协变量绝对标准化差异≤0.2。")
    add_para(doc, "预趋势不是用“不显著 p>0.05”来冒充相等，而是做等价门：处理前 24 小时的效应均值 95% CI 必须完全落在 ±0.03 内，同时趋势斜率 CI 必须落在 ±0.002/小时内。Sumy 的斜率约 −0.0099/小时，p=0.018，显然不满足，因此即使峰值很大也不能作为因果式主证据。")
    add_equation(doc, "max_deficit = max(0, baseline − nadir₀…W)", "W 为预先设定的缺口窗口；nadir 是最低可达性。")
    add_equation(doc, "AUC = Σₜ max(0, baseline − reachₜ) × 2 hours", "只累计低于基线的缺口；短而深与长而浅均可比较。")
    add_equation(doc, "t90 = first time after nadir with k consecutive cycles ≥ nadir + 0.9×max_deficit", "恢复时间从最低点而非攻击锚点开始；未在窗口内恢复则右删失。")

    doc.add_heading("14. 实验 C：重复性与预测为什么必须严格留出", level=2)
    add_para(doc, "所谓 ASN×Admin1 韧性指纹，至少需要同时满足“同一组在不同攻击中相对排序稳定”和“只用过去信息可以预测下一次攻击”。仅在一张热力图上看到某些组多次变红并不够，因为事件严重度和覆盖范围会制造共同排序。")
    add_bullets(doc, nums, [
        "重复性：只比较两个事件都被独立登记为暴露且都有结果的共同组；计算事件对 Spearman ρ，并按组 bootstrap；至少 3 个可比事件对。",
        "绝对一致性：ICC(1,1) 分解组间与组内变化；ICC=0 表示固定组效应没有从残差中脱颖而出。",
        "预测：按整个事件滚动留出；测试事件的当前结果绝不进入其特征；所有历史统计先 shift(1)。",
        "基线：M3 只记住 group 的历史均值；主模型 M4 Ridge 使用 ASN、Admin1、历史缺口、历史恢复、基线水平和样本规模。",
        "成功门：M4 相对 M3 的事件等权 MAE 至少改善 2%，至少 2/3 测试事件胜出，且事件置换 p<0.05。",
    ])
    add_para(doc, "结果中 M4 的表观 MAE 改善很大，但只有 4 个测试事件，置换 p=0.294；重复性 ρ 仅 0.179、ICC=0。方差分解进一步显示残差占 57.9%—86.2%。因此最稳妥的解释是“事件条件和未观测因素压倒固定 ASN×地区属性”，而不是“模型还需要更复杂”。继续堆深度模型只会增加过拟合。")

    doc.add_heading("15. 实验 D/E/F/G：为什么保留，但不让它们抢主线", level=2)
    add_table(doc, ["模块", "问题", "结果", "论文地位"], [
        ("D 恢复债务", "前次未恢复是否放大后续冲击", "预事件网络债务与 AUC 有关联；未来债务安慰剂接近但未显著；t90 不稳。官方 168h 暴露缺事件内变异。", "次要机制，写“association”，不写因果或电池耗尽。"),
        ("E 路径适应", "仍可达同一目标时 AS/ASGeo 是否变化", "1,250 组—事件仅 52 合格；6/52 ASGeo 经 BH-FDR 显著。", "条件于 reached targets；不能外推全网。"),
        ("F 外部验证", "内部异常是否与独立平台一致", "IODA 6 个事件仅 12/25 在 ±6h 起点一致；幅度 Spearman 很弱；其他外部来源有 2 个时间、2 个空间一致。", "三角验证与观测尺度差异，不是地面真值。"),
        ("G 州级 falsification", "同日执行州 vs 取消州是否形成差分", "07/24 仅 2 个周期；99 配对、7 ASN；平衡/预趋势失败；DID +0.0258，CI 跨 0。", "诚实的不可估计结果；说明队列—IP 映射的重要性。"),
        ("天气敏感性", "计划信号是否只是热浪代理", "排除官方热浪预警后 Δ 从 +0.007885 变 +0.005463，方向不反转；ERA5 连续覆盖为 0。", "部分敏感性；不能称天气调整完成。"),
    ], widths=[1.1, 1.55, 2.25, 1.35], font_size=8.1, caption="表 8　次要模块都围绕主问题，但不能反向升级主结论")

    doc.add_heading("16. 为什么别人可能不做、我们为什么仍能发表", level=2)
    add_para(doc, "别人不做并不意味着问题不重要。它通常来自四个门槛：需要事件发生前已经持续运行的主动测量；需要辨认扫描缺失与真实非响应；需要把不断更新、取消和恢复的乌克兰停电消息转成版本化机器表；还需要愿意把预期的正向故事设计成可能失败的留出检验。很多数据集只有网络曲线，没有电力标签；很多停电数据只有日级、国家级统计，没有 2 小时窗口；很多研究在看到异常后才挑事件，无法做真正前瞻冻结。")
    add_para(doc, "我们可以做，是因为当前数据同时满足长期性、端点分母、事件覆盖和证据登记四个条件。可以发表，不是因为技术算法前所未有，而是因为真实冲突环境提供了一个高价值的域迁移检验：从通知式、部分可规避的计划削减迁移到突发、破坏性能源攻击。结果表明这种迁移并不自动成立，这对灾害测量、基础设施风险研究和使用社会管理记录做弱监督的研究都有方法学价值。")

    doc.add_heading("17. 还需要做哪些闭环；是否值得增加实验", level=2)
    add_callout(doc, "对 v2.4 的行动建议", "不要重跑 v2.4，也不要再增加结果驱动阈值。现在最值得做的是材料闭环和写作；只有获得新的独立真值或新的未来事件，才开启 v2.5。", tone="orange")
    add_table(doc, ["优先级", "行动", "是否改变主结果", "目的/成功标准"], [
        ("P0 投稿前", "Telegram Desktop 导出官方频道原生 JSON+HTML+媒体，生成 selected_posts.csv、manifest 和 SHA-256", "否", "证明计划表、取消与恢复记录的原始文本和时间可追溯。"),
        ("P0 投稿前", "补齐可获取的官方网页原始响应；无法抓取时归档失败日志、研究捕获、URL 与哈希", "否", "避免把易变网页作为唯一证据。"),
        ("P1 建议", "下载 ERA5-Land 2m 温度，仅对冻结 exp_a_validation_long 做事先定义的负类残差化敏感性", "原则上否", "完成天气混杂审计；无论结果如何都报告。"),
        ("P1 建议", "整理匿名化复现包：事件表、周期质量、汇总传感器分母、所有结果表、图代码和环境", "否", "提高 artifact/reproducibility 可信度。"),
        ("v2.5 才做", "获得队列/馈线/地址级实际执行真值，或至少更多州级执行—取消自然对照", "可能", "区分“端点非传感器”与“L3 标签过粗”。"),
        ("v2.5 才做", "在尚未查看结果的新计划事件上逐事件学习 B2 并检验成员 Jaccard，再锁定未来攻击盲测", "可能", "直接检验供电敏感性是否为跨事件 IP 属性。"),
        ("扩展研究", "增加乌克兰内/邻国多测量源或 RIPE Atlas 视角", "可能", "分离源侧路径与目标侧掉线，改善空间归因。"),
    ], widths=[.8, 2.45, 1.05, 2.0], font_size=8.2, caption="表 9　新增工作的价值排序；P0/P1 不构成 v2.4 重跑")

    doc.add_heading("18. 指标与结果表的完整对应", level=2)
    add_table(doc, ["指标", "量化对象", "聚合/统计单位", "解释边界", "主要文件"], [
        ("pN / pP / S_lo", "IP 正常与计划窗口响应", "IP；Beta 后验", "候选敏感度，不是供电真值", "ip_sensor_scores_parts/*"),
        ("AUPRC", "计划正周期的排序质量", "完整 2h 周期；事件×Admin1 块", "受标签粒度和正类率影响", "exp_a_summary.csv"),
        ("ΔAUPRC", "B2 对 B1 的增量", "总体、事件、独立集群", "核心校准门", "exp_a_bootstrap_delta.csv"),
        ("normalized reach", "观察响应/期望响应", "/24×ASN×Admin1×周期", "可观测网络可达性", "sensor_event_panel/*.parquet"),
        ("immediate drop", "锚点前后短窗变化", "事件曲线", "可能受锚点不确定性影响", "exp_b_main_results.csv"),
        ("max deficit", "基线至最低点的深度", "事件/组—事件", "深度，不含持续性", "exp_b_main_results.csv"),
        ("deficit AUC", "低于基线的累计面积", "2h 离散积分", "深度×持续性", "exp_b_main_results.csv"),
        ("t90", "最低点后的 90% 恢复", "连续 k 周期", "可能右删失；不是电网恢复", "exp_b_main_results.csv"),
        ("pretrend equivalence", "处理前水平和斜率足够接近 0", "事件曲线", "失败则仅描述", "exp_b_estimand_results.csv"),
        ("Spearman / ICC", "跨事件排序/绝对一致性", "共同暴露组", "事件少时不确定性大", "exp_c_repeatability.csv"),
        ("event-equal MAE", "下一事件的预测误差", "滚动留出事件", "每个事件等权，防大事件支配", "exp_c_summary.csv"),
        ("JSD + BH-FDR", "同目标路径 AS/ASGeo 分布变化", "合格组—事件", "条件于成功到达目标", "exp_e_path_results.csv"),
    ], widths=[1.05, 1.45, 1.35, 1.45, 1.25], font_size=7.7, caption="表 10　从概念到代码产物的指标字典")

    doc.add_heading("第三部分　论文图像如何讲清核心结论", level=1)
    add_para(doc, "正文不应把所有生成过的图塞进去。推荐保留 6 张主图，其中前两张建立方法和事件证据，后四张直接回答 RQ1—RQ3；补充材料保留 4 张审稿人防御图。若投 13 页会议版本，可把图 1 与图 2 合并为“设计+时间轴”，形成 5 张主图。下面逐图解释现有有价值图像，不以旧编号多少决定价值。")

    add_report_figure_card(doc, "图 1", "从弱监督到可证伪结论的研究流程", FIG_M / "Fig1.png", [
        ("横轴/纵轴", "这是流程图，没有数值轴。阅读方向从左到右；垂直分层区分外部证据、端点筛选、留出攻击和组级推断。"),
        ("为什么这样画", "审稿人首先要知道哪些数据参与训练、哪里冻结、哪些只做外部验证；流程比十段方法文字更快揭示是否数据泄漏。"),
        ("应看到什么", "计划停电只产生候选监督；通过门才可选 B2，否则自动使用 B1；攻击与指纹均在冻结之后。"),
        ("说明的问题", "方法具备可证伪性与回退策略；“网络不可达≠确认停电”作为认识论边界。"),
        ("是否达到预期", "达到。它不声称正向结果，只解释为什么负结果仍是完整实验。"),
        ("主文位置", "引言贡献之后或方法开头。"),
        ("建议图注", "研究设计与信息隔离。外部计划停电用于弱监督，端点集合在攻击前冻结；第三方网络信号只做独立复核。"),
    ], "图 1　v2.4 的冻结分析链；每个箭头都有独立成功门。")

    add_report_figure_card(doc, "图 2", "计划停电表、战争事件表与可观测结果总览", FIG_M / "Fig2.png", [
        ("A 面板横轴", "日期：2024 年 6 月至 2025 年 1 月。"),
        ("A 面板纵轴", "事件角色泳道：计划训练、计划留出、攻击/盲测/压力测试。形状区分事件类型。"),
        ("B 面板横轴", "Europe/Kyiv 本地小时 0—24；每格对应计划表中的时间段。"),
        ("B 面板纵轴", "各计划停电事件；颜色表示最终登记的同时停电队列数 0—3。"),
        ("C 面板横轴", "B1 最大可达性缺口，越右表示事件中最深的网络响应下降。"),
        ("C 面板纵轴", "六个能源攻击及其国家/区域范围；灰色标记仅描述事件。"),
        ("为什么这样画", "把“监督来自哪些日子”“攻击是否严格在后”“计划是否是全天同强度”“哪些事件不能推断”一次交代清楚。"),
        ("核心结论", "计划监督并非简单的停电日标签；攻击效应高度异质；Sumy 即使缺口大也因预趋势失败不能进入主推断。"),
        ("是否达到预期", "达到设计透明度预期；不能单独证明因果或校准成功。"),
    ], "图 2　冻结事件登记、计划队列剂量和攻击缺口概览。")

    add_report_figure_card(doc, "图 3", "RQ1：计划停电校准门", FIG_C / "CoreFig1.png", [
        ("A 面板横轴", "B0、B1、B2 三种嵌套端点集合。"),
        ("A 面板纵轴", "留出计划周期分类的 AUPRC；越高表示正周期排名越靠前。"),
        ("B 面板横轴", "事件级 ΔAUPRC=B2−B1；0 线表示供电特异筛选无增量。"),
        ("B 面板纵轴", "各留出计划事件；07/28 的点远大于 8 月三次。"),
        ("为什么这样画", "A 显示整体性能，B 揭示平均值背后的事件支配；CI 和 p 值直接写在图中，避免仅凭柱高宣称成功。"),
        ("核心结论", "B1 相比 B0 有明显价值；B2 仅比 B1 高约 0.008，CI 跨 0、p=0.323，且增益由 07/28 主导。"),
        ("是否达到预期", "未达到原正向预期，但完整回答 RQ1：校准不成立。"),
        ("闭环作用", "触发冻结回退到 B1，后续不再称 B2 为已验证供电传感器。"),
    ], "图 3　稳定端点筛选有效，但供电特异 B2 没有通过预注册增量门。")

    add_report_figure_card(doc, "图 4", "RQ2：B2 是否泛化到留出攻击", FIG_C / "CoreFig2.png", [
        ("横轴", "每个事件的 B2−B1；左面板为最大缺口差，右面板为累计缺口 AUC 差。0 线表示两者相同。"),
        ("纵轴", "六个攻击事件；Sumy 灰色表示预趋势失败，计数总结只使用 5 个可推断事件。"),
        ("颜色/方向", "右侧意味着 B2 捕捉到更大缺口，左侧意味着 B2 反而更弱。"),
        ("为什么这样画", "直接画“增量”而非两套重复曲线，审稿人一眼能看出符号是否跨事件一致。"),
        ("核心结论", "最大缺口只有 3/5 为正，AUC 只有 2/5 为正；07/28 训练域的优势没有稳定迁移到攻击域。"),
        ("是否达到预期", "没有。它排除了“校准统计功效不够，但攻击中 B2 明显更好”的补救解释。"),
        ("闭环作用", "连接 RQ1 与攻击应用：B2 不仅校准未显著，跨域泛化也不一致。"),
    ], "图 4　供电特异候选端点在留出能源攻击中没有一致优于稳定端点。")

    add_report_figure_card(doc, "图 5", "RQ2：六轮攻击的事件时间动态", FIG_C / "CoreFig3.png", [
        ("横轴", "相对主电力锚点的小时，−24 至 +72；竖直 0 线是 outage_start_utc 或冻结回退锚点。"),
        ("纵轴", "B1 相对清洁基线/匹配对照的可达性缺口；正值表示低于基线，负值可表示短时超调或噪声。"),
        ("线与阴影", "线是事件级估计，阴影是按 /24 块 bootstrap 的 95% CI；标注峰值与 t90。"),
        ("为什么这样画", "峰值表不能显示冲击何时出现、是否先有趋势、恢复是否单调；事件曲线揭示持续性和锚点不确定性。"),
        ("核心结论", "8/26 和 12/13 深度较大；11/17/11/28 较浅但时间形态不同；12/25 很浅却恢复指标较长；深度与恢复不是同一维度。"),
        ("失败如何显示", "Sumy 整幅灰化并标“descriptive”，防止把最大缺口 0.170 当作最强因果结果。"),
        ("是否达到预期", "达到“部分攻击可量化”的预期，但不证明每次能源攻击都在网络上同样可见。"),
        ("闭环作用", "给出可观测韧性三元组：深度、累计损失、恢复时间。"),
    ], "图 5　留出能源攻击前后的 B1 可达性缺口与恢复。")

    add_report_figure_card(doc, "图 6", "RQ3：ASN×Admin1 指纹的重复性和预测", FIG_C / "CoreFig4.png", [
        ("A 面板横纵轴", "横轴和纵轴都是攻击事件；每个格子是两个事件中共同暴露组的 Spearman ρ。"),
        ("A 面板颜色", "正相关越深表示排序越一致；空白表示共同组不足，不能计算。"),
        ("B 面板横轴", "事件等权 MAE，越左越好。"),
        ("B 面板纵轴", "M3 组历史均值基线与 M4 历史 Ridge 主模型。"),
        ("为什么这样画", "热力图展示一致性不是普遍存在；MAE 面板把表观改善与置换显著性同时呈现。"),
        ("核心结论", "总体主 ρ≈0.179，ICC=0；M4 表观提升 41.7% 但 p=0.294。固定指纹不够稳定。"),
        ("是否达到预期", "未达到正向预期；这是核心负结果而非模型失败。"),
        ("闭环作用", "回答最终问题：当前证据不支持“可重复、可预测的 ASN—地区韧性指纹”。"),
    ], "图 6　跨事件排序弱且前瞻预测未通过置换门。")

    add_report_figure_card(doc, "图 S1", "校准的事件集群与高温敏感性", FIG_M / "FigS1.png", [
        ("A 横轴", "事件等权 ΔAUPRC；0 表示无增量。"),
        ("A 纵轴", "全样本、移除 August heat、移除 late-July 三种 leave-cluster-out。"),
        ("B 横轴/纵轴", "横轴仍为 ΔAUPRC；纵轴为全样本与排除官方热浪预警窗口。"),
        ("结论", "移除 late-July 后增益仅 0.0224；排除官方预警后 pooled 增益仍小，仅 0.00546。"),
        ("边界", "official warning 不是连续 ERA5 调整；该图只能说明预警窗口排除没有改变方向。"),
        ("位置", "补充材料；用于回应“单事件支配”和“热浪代理”质疑。"),
    ], "图 S1　校准增益对独立事件集群高度敏感，官方预警排除不改变主失败。")

    add_report_figure_card(doc, "图 S2", "州级执行—取消 falsification", FIG_M / "FigS2.png", [
        ("横轴", "Zaporizhzhia 执行相对 Volyn 取消的 DID 可达性效应及 95% CI；0 为无差异。"),
        ("纵轴", "2024-07-24 唯一冻结对照。"),
        ("结论", "点估计 +0.0258，CI [−0.0073,+0.0595]，方向也不符合清晰负效应；仅 2 个周期。"),
        ("为何不主文", "匹配平衡与预趋势均失败，且 Zaporizhzhia 只是部分队列，IP 无队列映射。"),
        ("价值", "展示作者尝试了更强的空间对照，并诚实判为 not estimable；为 v2.5 指向队列—地址真值。"),
    ], "图 S2　州级自然对照因暴露粒度和设计门失败而不可推断。")

    add_report_figure_card(doc, "图 S3", "IODA 外部核验", FIG_M / "FigS3.png", [
        ("A 横轴", "IODA 同时刻即时缺口与 24 小时最大缺口。"),
        ("A 纵轴", "六个攻击事件。"),
        ("B 横轴", "内部 B1 最大缺口。"),
        ("B 纵轴", "IODA 最大相对缺口；对角线为同量级。"),
        ("结论", "只有 12/25 在预设 ±6h 内出现 IODA onset；内部与 IODA 幅度排序相关很弱。"),
        ("解释", "两平台的目标、聚合层级和信号（BGP、主动探测、暗网）不同；不一致不是自动否定内部结果，但限制外部可复现性。"),
        ("为何不主文", "它不改变校准或指纹判定，只是观测尺度三角验证。"),
    ], "图 S3　IODA 与内部测量在少数事件上同向，但时序和幅度一致性有限。")

    add_report_figure_card(doc, "图 S4", "恢复债务与同目标存活路径变化", FIG_M / "FigS4.png", [
        ("A/B 横轴", "固定效应/组聚类模型的回归系数及 95% CI；分别预测累计缺口和 t90。"),
        ("A/B 纵轴", "真实的预事件网络债务与“未来债务”安慰剂。"),
        ("C 横轴", "每组—阶段要求的最小 traceroute 数 20/50/100。"),
        ("C 纵轴", "BH-FDR 后显著的 ASGeo 组—事件数。"),
        ("结论", "网络债务对累计缺口有条件关联；路径显著数为 6/52，在阈值 100 时为 8/33。"),
        ("边界", "恢复债务不等于官方停电小时；路径只覆盖仍可到达的相同目标，不能代表掉线端点。"),
        ("为何补充", "有机制价值，但样本选择和事后选择 exposure 的限制使其不能支撑主因果链。"),
    ], "图 S4　次要机制信号存在，但推断范围严格受可识别样本限制。")

    doc.add_heading("19. 哪些旧图或失败图不应进入正文", level=2)
    add_bullets(doc, nums, [
        "不要放 15 张流水账式结果图。覆盖率、月度漂移、所有 PR 点、所有模型散点可留在 artifact，不应占正文叙事空间。",
        "不要把 FigS2 的不合格 DID 画得像显著自然实验；标题和图注必须出现 not estimable。",
        "不要单独画高温预警色块并暗示已完成 ERA5；官方预警只能是排除式敏感性。",
        "不要把路径显著单元画成乌克兰地图大面积着色；分母只有 52 个合格单元，会产生虚假的全国覆盖感。",
        "不要只画 M4 比 M3 的两根柱而省略 p=0.294；这会让 41.7% 表观改善被误读成验证成功。",
        "不要用双 y 轴把攻击深度和 t90 强行叠加；它们意义不同，容易制造相关错觉。",
    ])

    doc.add_heading("20. 可直接用于论文的声明矩阵", level=2)
    add_table(doc, ["可写", "不可写", "原因"], [
        ("计划停电提供了方向一致但事件依赖的弱监督。", "计划停电成功校准了供电敏感 IP。", "B2 增量 CI 跨 0、p=0.323。"),
        ("B1 主动测量在部分留出攻击中观测到可达性缺口。", "我们远程测得了实际断电比例。", "Ping 不可达不是电力真值。"),
        ("B2 未在留出攻击中一致增强信号。", "攻击验证证明 B2 是供电特异传感器。", "5 个可推断事件中符号不一致。"),
        ("ASN×Admin1 响应表现出有限、事件依赖的相关。", "发现了稳定、可预测韧性指纹。", "ρ=0.179、ICC=0、预测 p=0.294。"),
        ("预事件网络恢复债务与累计缺口存在条件关联。", "备用电耗尽导致后续恢复变慢。", "没有电池/燃料观测，t90 结果不稳。"),
        ("存活且同目标的路径中有少量 ASGeo 分布变化。", "乌克兰网络普遍发生路由适应。", "只有 52/1250 合格，存在存活者条件。"),
    ], widths=[2.15, 2.15, 2.05], font_size=8.3, caption="表 11　写作口径是可信度的一部分")

    doc.add_heading("第四部分　投稿定位、导师答辩与执行路线", level=1)
    doc.add_heading("21. 论文标题与贡献应该怎样重写", level=2)
    add_para(doc, "推荐中文标题：〈计划停电不是端点断电真值：基于乌克兰长期主动测量的战时能源冲击可观测性与韧性指纹边界〉。推荐英文标题：Scheduled Outages Are Not Endpoint Ground Truth: Stress-testing Internet Sensors for Wartime Energy Shocks in Ukraine。")
    add_para(doc, "标题把负向结果变成中心，而不是在一个声称成功的系统标题下藏失败。贡献建议写成三点：第一，构建端点级零响应分母和版本化多源能源事件登记；第二，首次系统检验计划停电弱监督能否跨域迁移到留出战争能源攻击，并给出失败边界；第三，在不泄漏的滚动事件预测中检验 ASN×Admin1 指纹，并量化事件、组织/地区与残差方差。")

    doc.add_heading("22. 适合投什么", level=2)
    add_para(doc, "以下建议基于 2026-08-08 的公开页面。具体 2027 截止日期尚未正式公布时，不应自行臆测。")
    add_table(doc, ["目标", "适配度", "当前判断", "提交前必须强化"], [
        ("ACM IMC 2027 full paper", "最高但挑战大", "IMC 2026 full paper 两轮已截止；2026 poster 截止 2026-08-14，只适合两页反馈，不等于完整论文。[19][20]", "把稿件定位为弱监督边界；正面对比 IMC 2025 乌克兰全扫描和南非 poster；开放可审计 artifact。"),
        ("PAM 2027", "很匹配", "主动/被动测量、真实数据和方法学负结果契合；待官方 2027 CFP。[21]", "强调测量方法、标签粒度与跨域泛化。"),
        ("TMA 2027", "匹配", "适合流量/测量/分析与非常规挑战；待官方 CFP。[22]", "压缩故事，强化事件研究和复现。"),
        ("IEEE Transactions on Networking", "扩展后可投", "适合更完整的方法与长期研究。[23]", "最好补多视角或新真值/前瞻事件，形成比当前更强的普适性。"),
        ("Computer Networks", "当前版本较现实", "可容纳完整实证、负结果、详细鲁棒性和数据限制。[24]", "完成证据归档、ERA5 敏感性和英文稿。"),
    ], widths=[1.35, 1.0, 2.55, 1.45], font_size=8.2, caption="表 12　投稿路径不是按“名气”排序，而是按当前证据强度匹配")
    add_callout(doc, "推荐策略", "以 PAM/TMA 2027 或 Computer Networks 为现实主线；若目标是 IMC 2027 full paper，优先补强 artifact、队列/实际执行真值或多视角，而不是重调现有模型。IMC 2026 poster 只在你希望快速获得社区反馈时有意义。")

    doc.add_heading("23. 导师或审稿人最可能追问什么", level=2)
    add_table(doc, ["质疑", "直接回答", "证据"], [
        ("别人已经在 IMC 2025 测过乌克兰，为什么还要做？", "我们不重复“战争造成中断”，而是检验计划停电能否训练可迁移端点传感器，并在留出攻击与预测中证伪。", "严格训练/留出/回退链；B2 失败与 B1 可观测并存。"),
        ("Ping 掉线怎么能说是停电？", "不能，所以全文使用 observable reachability deficit；计划表只是 weak label。", "端点分母契约、外部三角验证、州级 falsification。"),
        ("为什么结果不显著还值得发？", "负结果揭示监督粒度边界，阻止把通知式调度直接当作实际供电真值。", "ΔAUPRC CI、集群支配、B2 攻击泛化不一致。"),
        ("41.7% 预测改善不是成功吗？", "不是。只有 4 个测试事件，事件置换 p=0.294，且 repeatability/ICC 失败。", "事件等权 MAE、置换、泄漏审计。"),
        ("单一 Frankfurt 视角是否足够？", "足以定义该视角下的可观测性，不足以断言目标侧物理原因。", "限制和未来多视角方案。"),
        ("为什么不继续调 B2？", "因为会用留出结果反向优化，破坏证伪性。", "冻结配置、自动 B1 回退、完整敏感性报告。"),
        ("天气和攻击同时发生怎么办？", "官方预警排除不改变方向；ERA5 尚待完成；因此不宣称去除了连续天气混杂。", "sens_official_heat_warning.csv 与缺失状态。"),
    ], widths=[1.85, 2.75, 1.75], font_size=8.1, caption="表 13　建议在答辩前熟悉的高概率问题")

    doc.add_heading("24. 投稿前执行清单", level=2)
    add_bullets(doc, nums, [
        "冻结并记录当前 run、代码 commit、Python 包、随机种子和配置哈希；不要覆盖 paper_v24_real_01。",
        "完成 Telegram 原生 JSON、HTML、媒体、selected_posts 与 SHA-256；核对消息编辑时间与最终调度版本。",
        "完成 ERA5 冻结敏感性或在主文明确标注缺失；不得把官方预警替代 ERA5。",
        "建立匿名 artifact：事件登记、质量报告、结果表、图生成脚本、最小示例数据、README 与数据伦理说明。",
        "正文控制为 5—6 张主图，补充 4 张；中英文图一一对应，正文语言版本只插同一种语言。",
        "在摘要、引言和结论三处使用一致口径：校准未验证、攻击可观测部分支持、固定指纹未验证。",
        "请独立合作者只根据图和图注复述结论；若其把 B2 说成成功或把缺口说成停电率，继续改图注。",
    ])

    doc.add_heading("附录 A　关键文件与复现入口", level=1)
    add_table(doc, ["内容", "路径（相对项目根）"], [
        ("冻结质量报告", "runs/paper_v24_real_01/results/tables/quality_report.json"),
        ("闭环报告", "runs/paper_v24_real_01/results/tables/closure_report.md"),
        ("校准结果", "runs/paper_v24_real_01/results/tables/exp_a_summary.csv"),
        ("攻击主结果", "runs/paper_v24_real_01/results/tables/exp_b_main_results.csv"),
        ("B2/B1 攻击敏感性", "runs/paper_v24_real_01/results/tables/exp_b_method_sensitivity.csv"),
        ("指纹总结", "runs/paper_v24_real_01/results/tables/exp_c_summary.csv"),
        ("恢复与路径", "runs/paper_v24_real_01/results/tables/exp_d_models.csv；exp_e_summary.csv"),
        ("核心中英文图", "runs/paper_v24_real_01/results/figures_submission_core/{zh,en}"),
        ("完整论文图", "runs/paper_v24_real_01/results/figures_manuscript/{zh,en}"),
        ("图再现命令", "MPLBACKEND=Agg MPLCONFIGDIR=.mplconfig PYTHONPATH=src python3 scripts/render_submission_core_figures.py --run-id paper_v24_real_01"),
    ], widths=[1.65, 4.7], font_size=8.4)

    doc.add_heading("附录 B　研究结论的最终一句话", level=1)
    add_callout(doc, "可作为答辩结束语", "在 2024-06 至 2025-01 的乌克兰长期 IP 主动测量中，计划停电记录能够提供弱且事件依赖的网络响应信号，但不足以校准一个经验证、能迁移到战争能源攻击的供电特异端点集合；冻结的稳定端点仍能量化部分攻击的可观测网络缺口与恢复，而 ASN×一级行政区响应在现有事件数量下不构成稳定可预测指纹。研究因此以清晰的适用边界而非预设的正向系统完成闭环。")

    add_reference_list(doc)
    doc.save(REPORT_PATH)
    return REPORT_PATH


def build_paper():
    doc = Document()
    configure_document(doc, "计划停电不是端点断电真值｜中文论文草稿")
    doc.core_properties.title = "计划停电不是端点断电真值：基于乌克兰长期主动测量的战时能源冲击可观测性与韧性指纹边界"
    doc.core_properties.subject = "中文论文正式草稿"
    doc.core_properties.author = "作者待补"
    add_cover(
        doc,
        "计划停电不是端点断电真值：\n基于乌克兰长期主动测量的战时能源冲击可观测性与韧性指纹边界",
        "Scheduled Outages Are Not Endpoint Ground Truth: Stress-testing Internet Sensors for Wartime Energy Shocks in Ukraine",
        "Chinese manuscript draft",
    )
    add_toc(doc, [
        "摘要与关键词",
        "1　引言",
        "2　背景与相关工作",
        "3　数据与证据登记",
        "4　方法",
        "5　结果：计划停电没有验证供电特异端点",
        "6　结果：留出能源攻击的可观测影响",
        "7　结果：ASN×Admin1 指纹未被验证",
        "8　条件性证据与失败分析",
        "9　讨论；10　局限、伦理与可复现性；11　结论",
        "附录与参考文献",
    ])
    nums = Numbering(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("作者：________　单位：________　通信作者：________")
    set_run_font(r, size=10, color=GRAY)

    doc.add_heading("摘要", level=1)
    add_para(doc, "战时能源攻击会影响互联网接入，但远程主动测量通常缺少端点级供电真值。本文检验一个具有实际吸引力的方案：利用精确登记的计划停电时间表作为弱监督，从长期 IP 级主动测量中筛选供电敏感端点，再将其冻结并迁移到未参与训练的能源攻击，最后判断 ASN×一级行政区响应能否形成可重复、可预测的韧性指纹。我们分析了 2024 年 6 月 22 日至 2025 年 1 月 9 日从 Frankfurt 每两小时对乌克兰目标执行的主动测量，共覆盖 3,042,264 个 IP；观测支持内 2,416 个周期中 96.11% 完整。我们构造三类嵌套端点集合：全部曾响应端点 B0、正常期稳定端点 B1 和计划停电期响应概率可信下降的候选 B2，并在独立计划事件上进行时间外验证。B2 的 AUPRC 为 0.308，B1 为 0.300，增量仅 0.007885（95% CI −0.025011 至 0.037959，事件置换 p=0.323），且增益主要由 2024 年 7 月 28 日事件驱动。按冻结规则，后续主分析回退到 361,871 个 B1 端点。B1 在若干留出能源攻击附近观测到可达性缺口，例如 8 月 26 日最大缺口 0.173、从最低点到 90% 恢复用时 100 小时；然而 B2 相对 B1 的攻击信号并不一致。跨事件 ASN×一级行政区累计缺口的主 Spearman 相关为 0.179，ICC 为 0；历史 Ridge 模型相对组均值基线的表观 MAE 改善为 41.7%，但置换 p=0.294。结果表明，计划停电可以提供方向一致但事件依赖的弱监督，却不能在当前标签粒度和事件数量下充当端点断电真值，也不足以建立可重复、可预测的固定韧性指纹。本文给出一套可审计的事件证据登记、零响应分母与跨域留出框架，并界定利用网络测量推断能源冲击时应遵守的证据边界。")
    p = doc.add_paragraph()
    r = p.add_run("关键词：")
    set_run_font(r, bold=True)
    p.add_run("互联网测量；能源攻击；计划停电；弱监督；网络韧性；乌克兰；事件研究")

    doc.add_heading("1　引言", level=1)
    add_para(doc, "能源系统与互联网接入之间存在直观依赖：基站、接入节点、用户终端和中继设施都需要供电。然而，在战争或大规模灾害中，研究者通常只能远程看到 BGP、流量、主动探测或暗网信号，而无法获得逐设施、逐地址的供电状态。把网络不可达直接解释为停电会混淆设备关机、路径故障、过滤、拥塞、测量丢失和用户行为。如何用公开能源记录校准一种可审计的网络代理，因此既是测量问题，也是因果归因问题。")
    add_para(doc, "乌克兰提供了一个重要而困难的研究环境。2024 年夏季存在大范围、有时按小时和队列发布的计划停电；秋冬又发生多轮能源基础设施攻击。与此同时，我们拥有从事件之前就持续运行的两小时 IP 级主动测量。这个组合使一种前瞻性设计成为可能：只用计划停电筛选端点，在完全留出的攻击中检验迁移，并在攻击之间检验地区—网络组织的稳定性。")
    add_para(doc, "这种设计也可能失败。计划停电是调度信号而非实际用户断电真值；公告可能临时取消或调整；队列无法映射到 IP；运营商和基站可能使用备用电；高温、战争损伤和网络故障会同时影响响应。因此，本文没有预设“供电传感器一定存在”，而是把校准、迁移和指纹分别设为必须通过的门。")
    add_para(doc, "本文围绕以下问题展开：精确登记的计划停电能否校准一组对供电变化敏感的互联网端点；这组端点能否在留出能源攻击上持续优于普通稳定端点；ASN×一级行政区响应是否跨事件重复并能被过去信息预测。")
    add_para(doc, "我们的主要发现是否定且具有边界意义。稳定端点筛选 B1 提升了计划周期检测能力，但进一步的供电特异筛选 B2 没有在预注册检验中显著优于 B1，也没有在留出攻击中一致增强信号。冻结 B1 仍可量化部分攻击的可观测可达性缺口与恢复，但同一 ASN×地区的响应受事件条件和未观测变化主导，未形成稳定预测指纹。")
    add_para(doc, "本文贡献如下：")
    add_bullets(doc, nums, [
        "构建一个端点级、零响应可见的长期主动测量数据契约，并把 3M 目标的稳定性与期望响应显式纳入分母；",
        "建立多源、版本化的计划停电与能源攻击证据登记，区分攻击开始、供电中断和第三方网络异常三个时钟；",
        "以严格时间留出和失败回退首次系统检验 scheduled-outage weak supervision 是否能够迁移到 wartime energy attacks；",
        "通过跨事件重复性、ICC 和无泄漏滚动预测检验 ASN×Admin1 韧性指纹，并报告不成立的统计边界；",
        "公开主结果与州级对照、天气、IODA、恢复债务和路径变化的限制，使负向结论可复核而非不可见。",
    ])

    add_figure(doc, FIG_M / "Fig1.png", "图 1　研究设计。计划停电只用于候选端点校准；端点集合在攻击前冻结；第三方网络信号只做复核。")

    doc.add_heading("2　背景与相关工作", level=1)
    doc.add_heading("2.1　主动测量与相关故障", level=2)
    add_para(doc, "Trinocular 等工作表明，低频主动探测与概率模型可以检测地址块级中断 [6]；相关故障研究进一步把 IP 异常聚合到地理或运营商层级 [13]。天气与日周期研究说明外部因素、终端作息和接入技术会系统性改变可达性 [3][14]。本文沿用“先识别可靠端点、再聚合事件效应”的思想，但将问题扩展到粗粒度能源弱标签和战争域迁移。")
    doc.add_heading("2.2　互联网与物理/供电中断", level=2)
    add_para(doc, "Down for Failure 使用住宅 IP 与县/州公用事业停电数据动态选择可靠端点 [12]，表明在更接近实际客户断电真值的条件下网络探测可以构成供电代理。Access Denied 从物理基础设施暴露评估接入风险 [4]，另有研究建模电网在互联网韧性中的作用 [17]。南非负荷削减研究发现计划削减对 NDT/BGP 的影响弱于非计划削减 [11]。这些工作提示监督粒度和冲击类型可能决定传感器能否迁移。")
    doc.add_heading("2.3　乌克兰战争中的互联网测量", level=2)
    add_para(doc, "已有研究使用 NDT、BGP、IXP、RIPE Atlas、IODA 和多源数据描述俄乌战争中的网络变化 [1][2][8][10][15][16][25]，近期政治学工作还将州—日级 IODA 中断用于研究战争行动与制度管理 [26]。最接近的 IMC 2025 工作同样分析两小时全地址块扫描并报告网络中断与停电日相关 [1][2]。本文不声称首次测量乌克兰网络中断；区别在于我们检验“计划停电能否训练端点传感器并迁移到攻击”，并进一步测试组级响应是否可重复、可预测。")

    doc.add_heading("3　数据与证据登记", level=1)
    doc.add_heading("3.1　主动测量", level=2)
    add_para(doc, "数据来自 Frankfurt 单一测量源，以约两小时为周期扫描固定目标集合。冻结观测支持为 2024-06-22 08:00 UTC 至 2025-01-09 14:00 UTC，共 2,416 个支持周期，其中 2,322 个通过完整采集门。目标集合含 3,042,264 个 IP；国家级 U2 分析保留 Ukraine+valid ASN，地区级 U3 进一步要求有效一级行政区。U2 含 2,145,724 个 IP、24,081 个 /24 与 1,542 个 ASN；U3 含 2,092,860 个 IP、23,307 个 /24 与 1,536 个 ASN。")
    add_para(doc, "在确认静态全扫描契约后，完整周期中某目标没有响应行被编码为非响应；若该周期的 Ping 采集文件缺失或不完整，则整个周期被标记为 acquisition gap，不作为零响应。IP 地理和 ASN 映射在端点层处理；一个 /24 中若存在多种映射，不用模态值覆盖全部端点。")
    doc.add_heading("3.2　计划停电与攻击事件", level=2)
    add_para(doc, "事件登记综合 Ukrenergo、州级配电运营商、政府/联合国/可靠报道与第三方网络平台。所有本地时间先按 Europe/Kyiv 解释，再转换到 UTC。对计划停电记录保存最终版本、队列数、零队列间隙、取消/恢复和 supersedes 关系。由于无法把 IP 映射到具体配电队列，主要监督等级是 L3 国家×时间，而不是 IP 实际断电。")
    add_para(doc, "对能源攻击分别登记 attack_start_utc、outage_start_utc 与 network_anomaly_start_utc。主电力估计使用独立的供电暴露范围和 outage anchor；第三方网络异常只构造 replication estimand。清洁基线在最早处理边界前 6 小时结束，并向前回看 168 小时。")
    add_figure(doc, FIG_M / "Fig2.png", "图 2　计划停电训练/留出事件、能源攻击时间轴、计划队列剂量与攻击最大缺口。灰色事件因预趋势失败仅作描述。")

    doc.add_heading("4　方法", level=1)
    doc.add_heading("4.1　端点集合与计划停电校准", level=2)
    add_para(doc, "设端点 i 在正常周期中响应 xN 次、暴露 nN 次，在训练计划周期中响应 xP 次、暴露 nP 次。采用 Jeffreys 先验得到后验均值 pN 与 pP，并定义 S=pN−pP 及其 95% 下界 S_lo。B0 包含训练期至少响应一次的端点；B1 进一步要求 pN 达到稳定阈值且至少有 24 个正常暴露周期；B2=B1∩{S_lo>0}。训练计划事件为 7 月 7 日和 7 月 20 日。")
    add_para(doc, "留出计划周期与同星期×同两小时时间槽的清洁负周期按 1:4 匹配。每种端点集合以响应缺口作为排序分数，主指标为 AUPRC。校准成功要求 B2−B1 的事件/地区块 bootstrap 95% CI 下界大于 0，并通过事件置换；发布性判断还要求至少两个独立事件集群。若失败，后续主方法冻结为 B1。")
    doc.add_heading("4.2　留出攻击事件研究", level=2)
    add_para(doc, "对每个 /24×ASN×Admin1×周期，令 expected_response_n=ΣpN，normalized_reach=responders/expected_response_n。国家事件比较事件曲线与清洁同 slot 自基线；区域事件将受影响地区的 /24 与未受影响地区中同 ASN、预事件特征最接近的 /24 配对。匹配后要求至少 30 对、协变量绝对标准化差异不超过 0.2。")
    add_para(doc, "事件曲线以电力锚点为 0。最大缺口为基线与事后最低点之差；累计缺口 AUC 为低于基线的正缺口按两小时间隔积分；t90 为从最低点开始连续恢复到最低点与基线差值的 90% 所需时间。处理前 24 小时同时通过水平±0.03和斜率±0.002/小时的 95% 等价区间，事件才进入推断。")
    doc.add_heading("4.3　重复性与前瞻预测", level=2)
    add_para(doc, "在两个事件中都被独立登记为暴露的共同 ASN×Admin1 组上计算 Spearman ρ 和组 bootstrap CI，并计算单向随机效应 ICC。预测以整个事件滚动留出：测试事件的所有结果均不可进入特征，组历史特征先滞后一事件。M3 为组历史均值基线，M4 为使用 ASN、Admin1 和历史/事前特征的 Ridge。主结果采用事件等权 MAE，并以事件置换评价相对改善。")
    doc.add_heading("4.4　次要与外部验证", level=2)
    add_para(doc, "我们另行检验前期网络恢复债务与下一事件的缺口、t90 关联；在同目标且追踪质量合格的存活路径上计算 AS/ASGeo 分布的 Jensen–Shannon divergence 并进行 BH-FDR；使用 IODA 和其他外部网络来源做事后三角核验；使用 7 月 24 日 Zaporizhzhia 执行、Volyn 取消的州级窗口做辅助 falsification。所有这些分析均不得选择 B2 或改变主锚点。")

    doc.add_heading("5　结果：计划停电没有验证供电特异端点", level=1)
    add_para(doc, "B0、B1、B2 分别包含 1,572,705、361,871 和 24,698 个 IP。留出计划周期中，B0 AUPRC 为 0.2074，B1 为 0.2997，B2 为 0.3076。稳定性筛选 B1 明显提高排序能力，但 B2 相对 B1 的增量仅 0.007885，95% CI 为 −0.025011 至 0.037959，事件置换 p=0.323。因此校准门失败，后续主分析使用 B1。")
    add_para(doc, "事件级增量显示强异质性：7 月 28 日为 +0.3937，8 月 19 日、20 日、21 日分别为 +0.0531、+0.0081、+0.0060；12 月 9 日处于攻击恢复混杂集群，发布性分析排除。对 4 个合格事件等权时 Δ 为 +0.1152；移除 late-July 集群后仅 +0.0224，说明总体方向受单一环境支配。")
    add_figure(doc, FIG_C / "CoreFig1.png", "图 3　计划停电校准。B2 仅比 B1 提高约 0.008，置信区间跨 0；事件级增益由 7 月 28 日主导。")

    doc.add_heading("6　结果：留出能源攻击的可观测影响", level=1)
    add_para(doc, "B1 在多轮能源攻击附近显示不同深度和持续时间的可达性缺口。8 月 26 日全国事件的最大缺口为 0.1733、累计缺口 2.834 小时、t90 为 100 小时；11 月 17 日区域事件分别为 0.0262、0.828 和 64 小时；11 月 28 日为 0.0282、0.259 和 2 小时；12 月 13 日为 0.1177、0.618 和 10 小时；12 月 25 日为 0.0102、0.348 和 66 小时。Sumy 9 月 17 日的表面最大缺口为 0.1696，但处理前斜率不满足等价门，故仅作描述。")
    add_para(doc, "U2 与 U3 目标宇宙的全国结果接近：例如 8 月 26 日最大缺口分别为 0.1733 和 0.1751，12 月 13 日为 0.11766 和 0.11764，12 月 25 日为 0.01022 和 0.01020。这说明国家结果对是否要求 Admin1 映射较稳健。")
    add_figure(doc, FIG_C / "CoreFig3.png", "图 4　六轮留出能源攻击的 B1 可达性缺口曲线。0 为独立登记的电力锚点；Sumy 灰色曲线仅作描述。")
    add_para(doc, "供电特异 B2 并没有在攻击域稳定优于 B1。在 5 个通过推断门的事件中，B2 的最大缺口仅 3 个更大，累计缺口仅 2 个更大；8 月 26 日和 12 月 13 日反而更弱。")
    add_figure(doc, FIG_C / "CoreFig2.png", "图 5　B2 相对 B1 的留出攻击增量。符号跨事件不一致，未形成稳定域迁移。")

    doc.add_heading("7　结果：ASN×Admin1 指纹未被验证", level=1)
    add_para(doc, "指纹分析包含 2,660 个组—事件，其中 641 个为攻击暴露组。累计缺口的主跨事件 Spearman ρ 为 0.1790，bootstrap CI 下界 −0.0115；ICC 为 0。虽然若干单独事件对有中等正相关，但相关并不普遍，且共同组数量差异大。")
    add_para(doc, "滚动留出产生 4 个测试事件。M4 Ridge 的事件等权 MAE 为 1.492，M3 组历史均值为 2.562，表观改善 41.7%，且四个测试事件均表面胜出；然而事件置换 p=0.294，未达到预注册 p<0.05。没有模型拟合失败，也没有时间泄漏警报，因此结论不是“代码没跑通”，而是“在事件数有限时，表观改善不能与偶然排序区分”。")
    add_para(doc, "方差分解支持这一解释。最大缺口中事件成分约 28.9%，ASN 6.9%，Admin1 1.9%，交互 4.4%，残差 57.9%；累计缺口残差 64.9%；t90 残差 86.2%。固定组身份无法解释大部分变异。")
    add_figure(doc, FIG_C / "CoreFig4.png", "图 6　ASN×Admin1 的跨事件重复性与滚动留出预测。相关较弱，表观 MAE 改善未通过置换检验。")

    doc.add_heading("8　条件性证据与失败分析", level=1)
    doc.add_heading("8.1　恢复债务", level=2)
    add_para(doc, "预注册官方 168 小时停电暴露在同一事件内缺少足够变化，固定效应模型不可识别。事后选择的 720 小时官方暴露对最大缺口、累计缺口和 t90 均不显著。由事件前 24 小时网络偏离构造的 pre_event_debt 与累计缺口呈正关联，但未来债务安慰剂接近显著边界，且该指标不是供电或电池观测。因此它只支持“尚未恢复的网络状态可能与后续累计损失有关”，不能证明能源储备耗尽机制。")
    doc.add_heading("8.2　路径变化", level=2)
    add_para(doc, "路径分析覆盖 1,250 个组—事件，但只有 52 个在同目标重叠、追踪数量和 AS/地理完整度上合格，约占 4.16%。其中 6 个 ASGeo 单元在 BH-FDR 后显著。该结果条件于事件期间仍能到达同一目标，因此只描述存活网络的转发分布变化；掉线端点和不可达路径不在样本中。")
    doc.add_heading("8.3　外部核验、天气与州级对照", level=2)
    add_para(doc, "IODA 结合 BGP、主动探测和网络望远镜等宏观信号 [7]。六个攻击均可取得 IODA 序列，但只有 12 月 25 日在预设 6 小时内出现 onset；内部与 IODA 最大缺口排序很弱。这表明不同测量平台看到不同目标与聚合尺度，外部一致性有限。官方热浪预警排除后 B2−B1 从 +0.007885 变为 +0.005463，方向不变但连续 ERA5 尚未完成。7 月 24 日州级执行—取消对照仅含两个测量周期，匹配平衡和预趋势均失败，因此不可推断。")

    doc.add_heading("9　讨论", level=1)
    doc.add_heading("9.1　为什么计划停电弱监督没有迁移", level=2)
    add_para(doc, "第一，公告描述的是队列或系统调度，不是具体 IP 所在馈线的实际执行；第二，计划停电可提前适应，用户和运营商可能启用电池、发电机或主动关机，产生与突发攻击不同的响应；第三，端点在线状态由接入网、用户设备和远端路径共同决定；第四，夏季热负荷、攻击恢复和调度版本增加事件间异质性；第五，独立计划事件集群仅两个，难以区分稳定属性与事件偶然性。")
    doc.add_heading("9.2　负结果的普适意义", level=2)
    add_para(doc, "公开管理记录经常被用作廉价监督标签，但“时间精确”不等于“个体暴露精确”。本文显示，即使计划表精确到小时和队列数，只要无法把暴露落到端点，模型就可能学习某个日期、天气或系统状态，而不是稳定的物理依赖。其他灾害、限电和基础设施研究应在主张传感器有效前要求事件集群外验证和处理域外迁移。")
    doc.add_heading("9.3　当前系统仍能做什么", level=2)
    add_para(doc, "B1 面板提供一种稳定的网络可观测性仪器：它能在独立登记的攻击时点估计从清洁基线偏离的深度、累计面积和恢复。但它不能区分供电、接入链路、路由或用户设备原因。适当的用途是事件响应、跨事件描述和发现需要进一步调查的地区/ASN，而不是宣布实际停电人口。")

    doc.add_heading("10　局限、伦理与可复现性", level=1)
    add_bullets(doc, nums, [
        "单一 Frankfurt 测量源限制因果定位；多源一致性尚未系统验证。",
        "Ping 不可达不是停电，主动测量可能受过滤、限速和终端行为影响。",
        "主要计划监督是 L3 国家×时间，缺少 IP—队列/馈线映射。",
        "连续 ERA5 天气协变量尚未归档；现有结果仅完成官方预警排除。",
        "原生 Telegram JSON 与个别官方网页原始响应需在 artifact 发布前补齐；统计结论不依赖其后调参。",
        "攻击事件数量有限，指纹检验的统计功效受限；不以复杂模型替代独立事件。",
        "路径结论存在存活者条件，只覆盖成功到达同一目标的 trace。",
        "发布数据应最小化可识别 IP 信息，优先提供聚合 /24、ASN×Admin1 和哈希化审计产物，并遵守扫描伦理与退出机制。",
    ])

    doc.add_heading("11　结论", level=1)
    add_para(doc, "本文检验了一个从计划停电弱监督出发、经过端点筛选、留出能源攻击和 ASN×Admin1 指纹预测的完整测量链。在乌克兰 2024 年 6 月至 2025 年 1 月的两小时 IP 级主动测量中，稳定端点筛选有效，但计划停电特异 B2 没有显著优于 B1，也没有在攻击域中稳定增强信号。冻结 B1 能够描述若干能源攻击的可观测可达性缺口与恢复，但组级响应的跨事件重复性和前瞻预测均未通过预注册门。计划停电因而可以作为弱、事件依赖的辅助监督，却不能在缺少端点级实际执行真值时被视为供电标签。这个负向边界为未来的电力—互联网联合测量明确了所需条件：更细的实际执行数据、更多独立事件和多测量源验证。")

    doc.add_heading("附录 A　补充图", level=1)
    add_figure(doc, FIG_M / "FigS1.png", "图 S1　独立事件集群与官方高温预警敏感性。校准增益受 late-July 事件支配；官方预警排除不改变主失败。")
    add_figure(doc, FIG_M / "FigS2.png", "图 S2　7 月 24 日州级执行—取消对照。仅两个周期且平衡/预趋势失败，结果不可推断。")
    add_figure(doc, FIG_M / "FigS3.png", "图 S3　IODA 外部核验。只有一个事件在 6 小时内达到预设时序一致性，内部/外部幅度排序很弱。")
    add_figure(doc, FIG_M / "FigS4.png", "图 S4　恢复债务和同目标存活路径的条件性证据。")

    doc.add_heading("附录 B　投稿前待办（不重跑 v2.4）", level=1)
    add_bullets(doc, nums, [
        "补齐 Telegram 原生 JSON/HTML/媒体和哈希清单；",
        "对冻结校准长表完成 ERA5-Land 连续温度敏感性，结果无论方向均报告；",
        "补齐或记录无法取得的官方网页原始响应；",
        "匿名化复现包并记录 run/config/code hash；",
        "将本文转为目标会议模板并压缩到 5—6 张正文图。",
    ])

    add_reference_list(doc)
    doc.save(PAPER_PATH)
    return PAPER_PATH


if __name__ == "__main__":
    report = build_report()
    paper = build_paper()
    print(report)
    print(paper)
